"""Agent 1 · Part 1 · Step 1 — DOCX document ingestion.

    In:  path to a .docx file
    Out: NormalizedDocument — ordered paragraphs (with style/formatting
         metadata) + tables (as row/column grids)

Mechanism: python-docx, walking the document body XML so paragraphs and tables
come back in true source order. Extraction is fully deterministic — no LLM call
belongs in this step (CLAUDE.md §29 assigns file ingestion and DOCX parsing to
deterministic code).

Scope boundary. This step captures what the document contains and stops there.
It does not detect sections (Step 2), parse questions (Step 4), or interpret any
routing, display or validation text (Part 2). Instruction text such as
"Show if: Q5 contains any touchpoint" is carried through verbatim as paragraph
or cell text; deciding what it means is explicitly not this step's job
(CLAUDE.md §7.1, §19).

Generalization. Nothing here keys off a section name, heading title, table
caption or column header seen in the sample corpus (CLAUDE.md §9, §10). The
reader walks whatever blocks the file contains, so a QRE with different
headings, extra tables or an unfamiliar layout still ingests without changes.
"""

from __future__ import annotations

from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .normalized_document import (
    SOURCE_FORMAT_DOCX,
    DocumentBlock,
    NormalizedDocument,
    NormalizedParagraph,
    NormalizedTable,
    SourceReference,
)


class DocxReadError(Exception):
    """Raised when a .docx file cannot be opened or read.

    Failing loudly is deliberate. CLAUDE.md §8 forbids silently producing
    incomplete extraction when document content cannot be read, and §40 states
    a visible failure is preferable to a silently incorrect result. Callers get
    an exception, never a half-populated NormalizedDocument.
    """


def _iter_block_items(document):
    """Yield the document body's paragraphs and tables in true source order.

    python-docx exposes `.paragraphs` and `.tables` as two separate lists, which
    loses the interleaving between them. The sample QRE interleaves heavily —
    a heading, then prose, then a questionnaire table, then more headings and
    tables — and that ordering is meaningful context. So walk the body XML
    directly and dispatch on element tag instead.

    Args:
        document: an open python-docx Document.

    Yields:
        Paragraph or Table objects, in document body order.
    """
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)
        # Other body children (sectPr, bookmarks, etc.) carry no QRE content
        # and are skipped. They hold no text, so nothing is lost per §16.


def _read_paragraph(item: Paragraph, index: int, document_name: str) -> NormalizedParagraph:
    """Convert one python-docx Paragraph into a NormalizedParagraph."""
    text = item.text

    # Direct run-level bold only. `run.bold` is tri-state in python-docx —
    # True (bold on), False (bold explicitly off), None (inherit from style) —
    # and None/False are both falsy, so any() gives exactly "some run is
    # explicitly bold". This is the mechanism the spec names for Step 3's
    # bold-label detection.
    is_bold = any(run.bold for run in item.runs)

    # A paragraph can carry no style object at all; treat that as the Word
    # default rather than raising or emitting None.
    style_name = item.style.name if item.style is not None else "Normal"

    # source_reference.text is left unset: `text` above IS this block's source
    # text, and storing it twice would double the JSON for no gain.
    return NormalizedParagraph(
        order_index=index,
        text=text,
        style_name=style_name,
        is_bold=is_bold,
        source_reference=SourceReference(document=document_name, order_index=index),
    )


def _read_table(item: Table, index: int, document_name: str) -> NormalizedTable:
    """Convert one python-docx Table into a NormalizedTable row/column grid."""
    rows = [[cell.text for cell in row.cells] for row in item.rows]

    return NormalizedTable(
        order_index=index,
        rows=rows,
        source_reference=SourceReference(document=document_name, order_index=index),
    )


def read_docx(path: str | Path) -> NormalizedDocument:
    """Ingest a .docx QRE into a NormalizedDocument.

    Args:
        path: path to a .docx file.

    Returns:
        NormalizedDocument with `blocks` holding every paragraph and table in
        document body order, each carrying provenance.

    Raises:
        DocxReadError: the file is missing, is not a .docx, or cannot be parsed.
    """
    path = Path(path)

    # Validate at the boundary so failures name the real cause. Without the
    # suffix check, handing this function a .pdf surfaces an opaque zip error.
    if not path.exists():
        raise DocxReadError(f"File not found: {path}")
    if path.suffix.lower() != ".docx":
        raise DocxReadError(
            f"Expected a .docx file, got '{path.suffix or 'no extension'}': {path}. "
            "Only DOCX ingestion is implemented at this stage."
        )

    try:
        document = docx.Document(str(path))
    except Exception as exc:  # python-docx raises varied errors for bad files
        raise DocxReadError(f"Could not read '{path}' as a .docx file: {exc}") from exc

    blocks: list[DocumentBlock] = []
    for index, item in enumerate(_iter_block_items(document)):
        if isinstance(item, Paragraph):
            blocks.append(_read_paragraph(item, index, path.name))
        else:
            blocks.append(_read_table(item, index, path.name))

    return NormalizedDocument(
        document_name=path.name,
        source_format=SOURCE_FORMAT_DOCX,
        blocks=blocks,
    )
