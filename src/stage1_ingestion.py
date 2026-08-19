"""Stage 1 — ingestion. DOCX to a document object in true body order.

Iterates the body XML so headings, paragraphs and tables stay interleaved;
`doc.paragraphs` and `doc.tables` are separate sequences and lose that ordering.

No LLM. Literal transcription only.
"""

from __future__ import annotations

import re
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from models import BlockKind, Paragraph, Stage1Document, Table

_HEADING_LEVEL = re.compile(r"^heading\s+(\d+)$", re.IGNORECASE)


class IngestionError(Exception):
    """File missing, wrong extension, or unreadable. Never a partial document."""


def _heading_level(style: str) -> int | None:
    match = _HEADING_LEVEL.match(style.strip())
    return int(match.group(1)) if match else None


def _iter_body(document) -> list[DocxParagraph | DocxTable]:
    items: list[DocxParagraph | DocxTable] = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            items.append(DocxParagraph(child, document))
        elif child.tag == qn("w:tbl"):
            items.append(DocxTable(child, document))
    return items


def run(path: str | Path) -> Stage1Document:
    path = Path(path)
    if not path.exists():
        raise IngestionError(f"File not found: {path}")
    if path.suffix.lower() != ".docx":
        raise IngestionError(f"Expected .docx, got '{path.suffix}': {path}")

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise IngestionError(f"Could not read '{path}': {exc}") from exc

    blocks: list[Paragraph | Table] = []
    for order, item in enumerate(_iter_body(document)):
        if isinstance(item, DocxParagraph):
            style = item.style.name if item.style is not None else "Normal"
            blocks.append(
                Paragraph(
                    kind=BlockKind.PARAGRAPH,
                    order=order,
                    text=item.text,
                    style=style,
                    is_bold=any(run.bold for run in item.runs),
                    heading_level=_heading_level(style),
                )
            )
        else:
            blocks.append(
                Table(
                    kind=BlockKind.TABLE,
                    order=order,
                    rows=[[cell.text for cell in row.cells] for row in item.rows],
                )
            )

    return Stage1Document(source=path.name, blocks=blocks)
