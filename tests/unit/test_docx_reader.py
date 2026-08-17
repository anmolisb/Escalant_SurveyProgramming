"""Regression check for Step 1 DOCX ingestion (CLAUDE.md tests/unit/).

Run directly: python tests/unit/test_docx_reader.py
Run via pytest: python -m pytest tests/unit/test_docx_reader.py -v
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

import docx

sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "src"))

from agents.qre_interpretation.ingestion.docx_reader import DocxReadError, read_docx
from agents.qre_interpretation.ingestion.normalized_document import (
    NormalizedParagraph,
    NormalizedTable,
)

FIXTURE = (
    Path(__file__).resolve().parents[2]
    / "fixtures"
    / "qre-samples"
    / "S01_campus_cafeteria_experience.docx"
)


def _make_sample_docx(path: Path) -> None:
    d = docx.Document()
    d.add_heading("Study specification", level=1)
    p = d.add_paragraph()
    p.add_run("Business objective:").bold = True
    p.add_run(" measure satisfaction.")
    d.add_paragraph("Plain instruction line.")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "ID"
    table.rows[0].cells[1].text = "Wording"
    table.rows[1].cells[0].text = "Q1"
    table.rows[1].cells[1].text = "Which channel?"
    d.save(str(path))


def test_reads_paragraphs_and_tables_in_order():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "sample.docx"
        _make_sample_docx(path)
        doc = read_docx(path)

        assert doc.source_format == "docx"
        assert doc.document_name == "sample.docx"

        kinds = [b.kind for b in doc.blocks]
        assert kinds == ["paragraph", "paragraph", "paragraph", "table"]

        heading = doc.blocks[0]
        assert isinstance(heading, NormalizedParagraph)
        assert heading.style_name == "Heading 1"
        assert heading.text == "Study specification"

        bold_para = doc.blocks[1]
        assert isinstance(bold_para, NormalizedParagraph)
        assert bold_para.is_bold is True
        assert "Business objective" in bold_para.text

        plain_para = doc.blocks[2]
        assert isinstance(plain_para, NormalizedParagraph)
        assert plain_para.is_bold is False

        table = doc.blocks[3]
        assert isinstance(table, NormalizedTable)
        assert table.rows == [["ID", "Wording"], ["Q1", "Which channel?"]]

        # order_index must match position, and blocks/paragraphs/tables views agree
        assert [b.order_index for b in doc.blocks] == [0, 1, 2, 3]
        assert doc.paragraphs == [heading, bold_para, plain_para]
        assert doc.tables == [table]


def test_source_reference_captures_provenance():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "sample.docx"
        _make_sample_docx(path)
        doc = read_docx(path)

        ref = doc.blocks[0].source_reference
        assert ref.document == "sample.docx"
        assert ref.order_index == 0
        # Blocks leave source_reference.text unset — the block's own `text`
        # field is its source text, so storing it twice would be redundant.
        assert ref.text is None
        assert doc.blocks[0].text == "Study specification"


def test_table_header_row_is_derived_not_stored():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "sample.docx"
        _make_sample_docx(path)
        doc = read_docx(path)

        table = doc.tables[0]
        assert table.header_row == ["ID", "Wording"]
        assert table.source_reference.text is None


def test_missing_file_raises_explicitly():
    try:
        read_docx("/nonexistent/path/does-not-exist.docx")
        raise AssertionError("expected DocxReadError")
    except DocxReadError:
        pass


def test_wrong_extension_raises_explicitly():
    """A non-.docx input must fail with a clear message, not an opaque zip error."""
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "not-a-word-file.pdf"
        path.write_bytes(b"%PDF-1.4 fake")
        try:
            read_docx(path)
            raise AssertionError("expected DocxReadError")
        except DocxReadError as exc:
            assert ".docx" in str(exc)


def test_unreadable_docx_raises_explicitly():
    """A corrupt file with the right extension must fail loudly (CLAUDE.md §8)."""
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "corrupt.docx"
        path.write_bytes(b"this is not a zip archive")
        try:
            read_docx(path)
            raise AssertionError("expected DocxReadError")
        except DocxReadError:
            pass


def test_real_fixture_if_present():
    """Sanity check against the actual sample QRE, when available locally."""
    if not FIXTURE.exists():
        return
    doc = read_docx(FIXTURE)
    assert len(doc.paragraphs) > 0
    assert len(doc.tables) >= 2  # questionnaire table + routing table
    # No content should vanish: every table row must be non-empty
    for table in doc.tables:
        assert all(row for row in table.rows)


if __name__ == "__main__":
    for _name, _fn in sorted(list(globals().items())):
        if _name.startswith("test_") and callable(_fn):
            _fn()
    print("OK: all docx_reader checks passed")
