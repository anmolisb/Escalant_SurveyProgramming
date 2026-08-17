"""Regression check for Step 4 question extraction.

Run directly: python3 tests/unit/test_question_parser.py
Run via pytest: python3 -m pytest tests/unit/test_question_parser.py -v
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

import docx

sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "src"))

from agents.qre_interpretation.extraction.question_parser import (
    DEFAULT_COLUMN_ALIASES,
    parse_questions,
)
from agents.qre_interpretation.extraction.section_detector import detect_sections
from agents.qre_interpretation.ingestion.docx_reader import read_docx

FIXTURE = (
    Path(__file__).resolve().parents[2]
    / "fixtures"
    / "qre-samples"
    / "S01_campus_cafeteria_experience.docx"
)

SPEC_FIELDS = ("id", "wording", "type", "options_raw", "display_validation_raw")


def _questionnaire_docx(path: Path, rows: list[list[str]], heading: str = "Questionnaire"):
    """Build a .docx with one heading and one table under it."""
    d = docx.Document()
    h = d.add_paragraph(heading)
    h.style = d.styles["Heading 1"]
    table = d.add_table(rows=len(rows), cols=len(rows[0]))
    for r, row in enumerate(rows):
        for c, cell in enumerate(row):
            table.rows[r].cells[c].text = cell
    d.save(str(path))


def _extract(tmp: str, rows: list[list[str]], **kwargs):
    path = Path(tmp) / "q.docx"
    _questionnaire_docx(path, rows)
    return parse_questions(detect_sections(read_docx(path)), **kwargs)


STANDARD = [
    ["ID", "Wording / instruction", "Type", "Options / scale", "Display / validation"],
    ["Q1", "How satisfied?", "single", "Yes; No", "Always show"],
    ["Q2", "Why?", "text", "—", "Show if: Q1 == 'Yes'"],
]


# ---------------------------------------------------------------------------
# Output contract — the five specified fields
# ---------------------------------------------------------------------------


def test_produces_the_five_specified_fields():
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, STANDARD)
        assert len(qe.questions) == 2
        q1 = qe.questions[0]
        for name in SPEC_FIELDS:
            assert hasattr(q1, name), f"missing specified field: {name}"
        assert q1.id == "Q1"
        assert q1.wording == "How satisfied?"
        assert q1.type == "single"
        assert q1.options_raw == "Yes; No"
        assert q1.display_validation_raw == "Always show"


def test_row_order_and_row_index_preserved():
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, STANDARD)
        assert [q.id for q in qe.questions] == ["Q1", "Q2"]
        assert [q.row_index for q in qe.questions] == [1, 2]


def test_column_mapping_is_reported():
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, STANDARD)
        assert qe.column_mapping["id"] == "ID"
        assert qe.column_mapping["options_raw"] == "Options / scale"
        assert qe.unmapped_headers == []


# ---------------------------------------------------------------------------
# The "as-is, no interpretation" boundary
# ---------------------------------------------------------------------------


def test_options_are_not_split():
    """'Yes; No' must stay one string — splitting belongs to Step 5."""
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, STANDARD)
        options = qe.questions[0].options_raw
        assert isinstance(options, str)
        assert options == "Yes; No"


def test_validation_json_is_not_parsed():
    rows = [
        STANDARD[0],
        ["Q1", "Pick some", "multi", "A; B", 'Validate: {"min_selections": 1}'],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        raw = qe.questions[0].display_validation_raw
        assert isinstance(raw, str)
        assert raw == 'Validate: {"min_selections": 1}'


def test_multiline_display_cell_keeps_its_newline():
    """Two instructions in one cell stay one raw string; Step 5 splits them."""
    cell = "Show if: Q1 == 'Yes'\nValidate: {\"min_length\": 10}"
    rows = [STANDARD[0], ["Q2", "Describe", "text", "—", cell]]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert qe.questions[0].display_validation_raw == cell
        assert "\n" in qe.questions[0].display_validation_raw


def test_em_dash_options_not_interpreted_as_empty():
    """Deciding '—' means 'no options' is interpretation, not extraction."""
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, STANDARD)
        assert qe.questions[1].options_raw == "—"
        assert qe.questions[1].empty_fields == []


def test_type_is_not_mapped_to_a_platform_type():
    rows = [STANDARD[0], ["Q1", "W", "single_select_radio", "A; B", "Always show"]]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert qe.questions[0].type == "single_select_radio"


# ---------------------------------------------------------------------------
# Dynamic column matching (CLAUDE.md §9, §10)
# ---------------------------------------------------------------------------


def test_alternative_column_names_are_matched():
    """A QRE naming its columns differently must still parse."""
    rows = [
        ["Question No", "Question Text", "Answer Type", "Codeframe", "Base"],
        ["1", "How old are you?", "numeric", "18-99", "Always show"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert len(qe.questions) == 1
        q = qe.questions[0]
        assert (q.id, q.wording, q.type) == ("1", "How old are you?", "numeric")
        assert q.options_raw == "18-99"
        assert q.display_validation_raw == "Always show"


def test_column_order_does_not_matter():
    rows = [
        ["Type", "Options / scale", "ID", "Display / validation", "Wording"],
        ["single", "Yes; No", "Q9", "Always show", "Agree?"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        q = qe.questions[0]
        assert (q.id, q.wording, q.type, q.options_raw) == (
            "Q9",
            "Agree?",
            "single",
            "Yes; No",
        )


def test_custom_column_aliases_override_default():
    rows = [["Kennung", "Fragetext"], ["Q1", "Wie alt?"]]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(
            tmp,
            rows,
            column_aliases={"id": ("kennung",), "wording": ("fragetext",)},
        )
        assert qe.questions[0].id == "Q1"
        assert qe.questions[0].wording == "Wie alt?"


def test_partial_header_phrasing_matches():
    """Regression: exact whole-string alias matching failed on real variations.

    'Q No' and 'Base / Validation' both went unmatched, so a QRE phrasing its
    columns slightly differently than the alias list extracted zero questions.
    Matching is now token-subsequence based.
    """
    rows = [
        ["Q No", "Question Text", "Answer Type", "Codeframe", "Base / Validation"],
        ["SC1", "Do you have a card?", "single", "1=Yes; 2=No", "ask all"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        q = qe.questions[0]
        assert q.id == "SC1"
        assert q.wording == "Do you have a card?"
        assert q.type == "single"
        assert q.options_raw == "1=Yes; 2=No"
        assert q.display_validation_raw == "ask all"


def test_best_matching_column_wins_not_the_leftmost():
    """A weak early match must not claim a role a later column fits better.

    'Scripter notes' precedes 'Base / Validation'; the latter is the real
    validation column, so roles are assigned by score, not document order.
    """
    rows = [
        ["Scripter notes", "Q No", "Question Text", "Base / Validation"],
        ["randomise list", "SC1", "Do you have a card?", "ask all"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert qe.column_mapping["display_validation_raw"] == "Base / Validation"
        assert qe.unmapped_headers == ["Scripter notes"]
        # ...and the scripter note is still preserved, not dropped
        assert [e.value for e in qe.questions[0].extra_columns] == ["randomise list"]


def test_incidental_keyword_does_not_claim_a_role():
    """'No. of visits' mentions 'no' but is not an id column.

    Coverage scoring rejects it: one matched token out of three is below the
    threshold, so the column goes to review rather than being mis-assigned.
    """
    rows = [
        ["ID", "Question Text", "No. of visits allowed"],
        ["Q1", "How many?", "3"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert qe.column_mapping["id"] == "ID"
        assert "No. of visits allowed" in qe.unmapped_headers


def test_token_matching_does_not_match_inside_words():
    """'id' must not match inside 'Validation', nor 'no' inside 'Notes'."""
    from agents.qre_interpretation.extraction.label_matching import score_label

    assert "id" not in score_label("Validation", DEFAULT_COLUMN_ALIASES)
    assert "id" not in score_label("Guidance", DEFAULT_COLUMN_ALIASES)


def test_header_found_below_a_leading_title_row():
    rows = [
        ["Main questionnaire", "", "", "", ""],
        STANDARD[0],
        ["Q1", "How satisfied?", "single", "Yes; No", "Always show"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert len(qe.questions) == 1
        assert qe.questions[0].id == "Q1"


# ---------------------------------------------------------------------------
# Nothing discarded, problems flagged (CLAUDE.md §16, §30, §45)
# ---------------------------------------------------------------------------


def test_unmapped_column_content_is_preserved_and_flagged():
    rows = [
        ["ID", "Wording", "Translator notes"],
        ["Q1", "How satisfied?", "keep formal register"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        q = qe.questions[0]
        assert [(e.header, e.value) for e in q.extra_columns] == [
            ("Translator notes", "keep formal register")
        ]
        assert "Translator notes" in qe.unmapped_headers
        assert any(r.reason == "unmapped_columns_preserved" for r in qe.review_queue)


def test_blank_header_columns_do_not_collapse_onto_each_other():
    """Regression: two blank headers once collapsed to one key, losing content.

    `extra_columns` was a {header_text: value} dict, so every column under a
    blank header mapped to the key "" and all but the last were silently dropped
    — a §16 violation. Now keyed by column index.
    """
    rows = [
        ["ID", "Wording", "", ""],
        ["Q1", "How satisfied?", "first extra", "second extra"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        values = {e.value for e in qe.questions[0].extra_columns}
        assert values == {"first extra", "second extra"}
        assert [e.column_index for e in qe.questions[0].extra_columns] == [2, 3]
        # A blank header is not a reportable column name
        assert qe.unmapped_headers == []


def test_duplicate_unmapped_header_names_both_preserved():
    """Regression: same-named unmapped columns collapsed in the old dict."""
    rows = [
        ["ID", "Wording", "Reviewer", "Reviewer"],
        ["Q1", "How satisfied?", "alice", "bob"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert {e.value for e in qe.questions[0].extra_columns} == {"alice", "bob"}


def test_empty_optional_field_is_not_reported_as_a_defect():
    """Regression: an empty options cell on a text question was a false positive.

    CLAUDE.md §18 requires applicability and requirement to stay distinct. A free-
    text question has no options; reporting that as malformed buries real findings.
    """
    rows = [
        STANDARD[0],
        ["Q1", "Describe the issue", "text", "", "Always show"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert qe.questions[0].empty_fields == ["options_raw"]
        assert qe.review_queue == []


def test_empty_required_field_is_still_reported():
    """The counterpart: an unusable row must still be flagged."""
    rows = [STANDARD[0], ["", "Describe the issue", "text", "x", "y"]]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert any(
            r.reason == "required_question_fields_empty" for r in qe.review_queue
        )


def test_missing_required_column_blocks_extraction_with_a_report():
    """No id column means rows cannot be identified — report, do not invent."""
    rows = [["Wording", "Type"], ["How satisfied?", "single"]]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert qe.questions == []
        assert any(
            r.reason == "required_columns_not_identified" for r in qe.review_queue
        )


def test_missing_optional_column_is_flagged_and_field_left_empty():
    rows = [["ID", "Wording"], ["Q1", "How satisfied?"]]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        q = qe.questions[0]
        assert q.type == ""
        assert q.options_raw == ""
        assert "type" in q.empty_fields
        reasons = [r.reason for r in qe.review_queue]
        assert reasons.count("optional_column_not_identified") == 3


def test_empty_cell_is_flagged_not_filled_in():
    rows = [STANDARD[0], ["Q1", "How satisfied?", "", "Yes; No", "Always show"]]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert qe.questions[0].type == ""
        assert qe.questions[0].empty_fields == ["type"]
        # An empty optional field is a fact, not a defect — no review item (§18)
        assert qe.review_queue == []


def test_blank_spacer_rows_are_skipped():
    rows = [
        STANDARD[0],
        ["Q1", "How satisfied?", "single", "Yes; No", "Always show"],
        ["", "", "", "", ""],
        ["Q2", "Why?", "text", "—", "Always show"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert [q.id for q in qe.questions] == ["Q1", "Q2"]


def test_duplicate_question_ids_are_flagged():
    rows = [
        STANDARD[0],
        ["Q1", "First", "single", "Yes; No", "Always show"],
        ["Q1", "Second", "single", "Yes; No", "Always show"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        # Both rows are kept — dropping one would lose content
        assert len(qe.questions) == 2
        assert any(r.reason == "duplicate_question_id" for r in qe.review_queue)


def test_question_table_found_even_under_the_wrong_heading():
    """A question table under a non-questionnaire heading must still be found.

    Step 2 mislabelling or failing to classify one heading must not take Step 4
    down with it. A table with both an id and a wording column is question-shaped
    whatever section it sits in — found, and reported as found out of place.
    """
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "q.docx"
        _questionnaire_docx(path, STANDARD, heading="Routing and termination")
        qe = parse_questions(detect_sections(read_docx(path)))
        assert [q.id for q in qe.questions] == ["Q1", "Q2"]
        assert any(
            r.reason == "question_table_found_outside_named_section"
            for r in qe.review_queue
        )


def test_question_table_found_under_an_unclassifiable_heading():
    """The cascading-failure case: Step 2 could not name the section at all."""
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "q.docx"
        _questionnaire_docx(path, STANDARD, heading="Instrument Content Matrix")
        qe = parse_questions(detect_sections(read_docx(path)))
        assert [q.id for q in qe.questions] == ["Q1", "Q2"]
        origin = next(
            r for r in qe.review_queue
            if r.reason == "question_table_found_outside_named_section"
        )
        assert "Instrument Content Matrix" in origin.element


def test_non_question_tables_are_not_mistaken_for_the_questionnaire():
    """The fallback must not grab a routing or acceptance-test table.

    Requiring BOTH id and wording is what keeps them out: a routing table maps
    neither, an acceptance-test table maps id but not wording.
    """
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "q.docx"
        d = docx.Document()
        h = d.add_paragraph("Some Unknown Section")
        h.style = d.styles["Heading 1"]
        for rows in (
            [["Rule", "Condition", "Action", "Destination"],
             ["R1", "S1 == 'No'", "terminate", "TERM"]],
            [["ID", "Purpose", "Key inputs", "Expected outcome"],
             ["T1", "screenout", '{"S1": "No"}', '{"end": "TERM"}']],
        ):
            t = d.add_table(rows=len(rows), cols=len(rows[0]))
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    t.rows[r].cells[c].text = cell
            d.add_paragraph("")
        d.save(str(path))

        qe = parse_questions(detect_sections(read_docx(path)))
        assert qe.questions == []
        assert any(
            r.reason == "questionnaire_section_not_found" for r in qe.review_queue
        )


def test_questionnaire_section_without_a_table_is_reported():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "q.docx"
        d = docx.Document()
        h = d.add_paragraph("Questionnaire")
        h.style = d.styles["Heading 1"]
        d.add_paragraph("Q1. How satisfied are you? (single: Yes/No)")
        d.save(str(path))

        qe = parse_questions(detect_sections(read_docx(path)))
        assert qe.questions == []
        assert any(
            r.reason == "no_table_in_questionnaire_section" for r in qe.review_queue
        )


def test_short_row_yields_empty_fields_not_an_error():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "q.docx"
        # Build a table then blank trailing cells to simulate a ragged row.
        _questionnaire_docx(
            path, [STANDARD[0], ["Q1", "How satisfied?", "single", "", ""]]
        )
        qe = parse_questions(detect_sections(read_docx(path)))
        q = qe.questions[0]
        assert q.options_raw == ""
        assert q.display_validation_raw == ""
        assert set(q.empty_fields) == {"options_raw", "display_validation_raw"}
        assert qe.review_queue == []


def test_questions_split_across_two_tables_are_concatenated():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "q.docx"
        d = docx.Document()
        h = d.add_paragraph("Questionnaire")
        h.style = d.styles["Heading 1"]
        for rows in ([STANDARD[0], STANDARD[1]], [STANDARD[0], STANDARD[2]]):
            t = d.add_table(rows=len(rows), cols=len(rows[0]))
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    t.rows[r].cells[c].text = cell
            d.add_paragraph("")
        d.save(str(path))

        qe = parse_questions(detect_sections(read_docx(path)))
        assert [q.id for q in qe.questions] == ["Q1", "Q2"]
        assert len(qe.source_tables) == 2


# ---------------------------------------------------------------------------
# LLM column classifier fallback (stubbed — no network)
# ---------------------------------------------------------------------------


def test_classifier_places_columns_the_vocabulary_cannot():
    """An unfamiliar header must reach the classifier and be usable."""
    calls = []

    def classifier(header, open_roles):
        calls.append(header)
        return {"Verbatim prompt shown": "wording", "Tracking stub": "id"}.get(header)

    rows = [
        ["Tracking stub", "Verbatim prompt shown", "Type"],
        ["SC1", "Do you have a card?", "single"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows, classifier=classifier)
        assert [q.id for q in qe.questions] == ["SC1"]
        assert qe.questions[0].wording == "Do you have a card?"
        # 'Type' resolved deterministically, so it never reached the model
        assert "Type" not in calls


def test_classifier_not_consulted_when_vocabulary_suffices():
    calls = []

    def classifier(header, open_roles):
        calls.append(header)
        return None

    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, STANDARD, classifier=classifier)
        assert len(qe.questions) == 2
        assert calls == [], "all five columns matched deterministically"


def test_classifier_cannot_return_a_role_outside_the_vocabulary():
    """§17: a classifier must not widen the contract.

    The invented role must never appear in the mapping. Extraction may still
    succeed via value-shape inference — that is the intended layering — but the
    inference is reported, and nothing is credited to the bogus role.
    """
    rows = [["Tracking stub", "Verbatim prompt shown"], ["SC1", "Do you have a card?"]]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(
            tmp, rows, classifier=lambda header, open_roles: "invented_role"
        )
        assert "invented_role" not in qe.column_mapping
        for question in qe.questions:
            assert not hasattr(question, "invented_role")
        if qe.questions:
            assert any(
                r.reason == "column_role_inferred_from_values"
                for r in qe.review_queue
            )


def test_classifier_cannot_steal_a_deterministically_matched_role():
    """Only unfilled roles are offered, so a real match cannot be overridden."""
    offered: list[list[str]] = []

    def classifier(header, open_roles):
        offered.append(list(open_roles))
        return None

    rows = [["ID", "Question Text", "Scripter stub"], ["Q1", "How old?", "x"]]
    with tempfile.TemporaryDirectory() as tmp:
        _extract(tmp, rows, classifier=classifier)
        assert offered, "classifier should have been consulted for 'Scripter stub'"
        for roles in offered:
            assert "id" not in roles and "wording" not in roles


def test_classifier_returning_none_leaves_column_unmapped_and_preserved():
    rows = [
        ["ID", "Question Text", "Revision history"],
        ["Q1", "How old?", "v2 2026-03-01"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows, classifier=lambda header, open_roles: None)
        assert "Revision history" in qe.unmapped_headers
        assert [e.value for e in qe.questions[0].extra_columns] == ["v2 2026-03-01"]


# ---------------------------------------------------------------------------
# Value-shape inference (deterministic, no model, no data sent anywhere)
# ---------------------------------------------------------------------------


def test_id_and_wording_inferred_from_value_shape():
    """Wholly unfamiliar headers still parse, from the shape of the values.

    'Marker' says nothing about being an id, but A_01/A_02/A_03 do: short,
    all distinct, no internal spaces. The wording column is the longest
    sentence-like one.
    """
    rows = [
        ["Marker", "Verbatim shown to participant", "Localisation memo"],
        ["A_01", "In which year were you born?", "keep numerals"],
        ["A_02", "Which of these have you used?", "translate list"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)  # no classifier at all
        assert [q.id for q in qe.questions] == ["A_01", "A_02"]
        assert qe.questions[0].wording == "In which year were you born?"
        assert qe.column_mapping["id"] == "Marker"
        assert any(
            r.reason == "column_role_inferred_from_values" for r in qe.review_queue
        )


def test_inference_does_not_override_a_header_match():
    """A column named outright must win over any value-shape guess."""
    rows = [
        ["Notes stub", "ID", "Question Text"],
        ["x1", "Q1", "How old are you?"],
        ["x2", "Q2", "Where do you live?"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert qe.column_mapping["id"] == "ID"
        assert qe.column_mapping["wording"] == "Question Text"
        assert not any(
            r.reason == "column_role_inferred_from_values" for r in qe.review_queue
        )


def test_inference_declines_when_no_column_is_id_shaped():
    """Repeated non-unique values are not an id column; do not force a match."""
    rows = [
        ["Alpha", "Beta"],
        ["same", "also repeated text here"],
        ["same", "also repeated text here"],
    ]
    with tempfile.TemporaryDirectory() as tmp:
        qe = _extract(tmp, rows)
        assert qe.questions == []
        assert any(
            r.reason in ("required_columns_not_identified",
                         "questionnaire_section_not_found")
            for r in qe.review_queue
        )


# ---------------------------------------------------------------------------
# Real fixture
# ---------------------------------------------------------------------------


def test_real_fixture_if_present():
    if not FIXTURE.exists():
        return
    qe = parse_questions(detect_sections(read_docx(FIXTURE)))

    assert [q.id for q in qe.questions] == [
        "S1", "S2", "Q1", "Q2", "Q3", "Q4", "Q5", "Q6", "Q7", "Q8",
    ]
    assert qe.unmapped_headers == []
    assert qe.review_queue == []
    assert set(qe.column_mapping) == set(SPEC_FIELDS)

    byid = qe.by_id()
    # Raw content survives exactly as written
    assert byid["Q2"].options_raw == (
        "1 - Very poor; 2 - Poor; 3 - Fair; 4 - Good; 5 - Excellent"
    )
    assert byid["Q4"].display_validation_raw == 'Validate: {"min_selections": 1}'
    assert byid["Q6"].options_raw == "—"
    assert "\n" in byid["Q6"].display_validation_raw
    assert byid["Q6"].display_validation_raw.startswith("Show if: Q5 == 'Yes'")
    # Types are whatever the document said
    assert {q.type for q in qe.questions} == {"single", "multi", "text"}
    # Provenance points at the questionnaire table block
    assert all(q.source_reference.order_index == 10 for q in qe.questions)


if __name__ == "__main__":
    for _name, _fn in sorted(list(globals().items())):
        if _name.startswith("test_") and callable(_fn):
            _fn()
    print(
        f"OK: all question_parser checks passed "
        f"({len(DEFAULT_COLUMN_ALIASES)} column roles)"
    )
