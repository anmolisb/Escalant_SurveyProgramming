"""Regression check for Step 2 dynamic section detection.

Run directly: python3 tests/unit/test_section_detector.py
Run via pytest: python3 -m pytest tests/unit/test_section_detector.py -v
"""

from __future__ import annotations

import os

# Keep this file hermetic when run directly (pytest gets the same via
# tests/conftest.py). The pipeline steps default to the configured LLM, so
# without this a standalone run would make real API calls.
os.environ["LLM_PROVIDER"] = "none"
os.environ["GROQ_API_KEY"] = ""

import sys
import tempfile
from pathlib import Path

import docx

sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "src"))

from agents.qre_interpretation.extraction.section_detector import (
    DEFAULT_SECTION_TYPES,
    detect_sections,
)
from agents.qre_interpretation.extraction.sectioned_document import (
    LABEL_DERIVED,
    LABEL_EXTRACTED,
    LABEL_INFERRED,
    LABEL_UNKNOWN,
    PREAMBLE_LABEL,
)
from agents.qre_interpretation.ingestion.docx_reader import read_docx

FIXTURE = (
    Path(__file__).resolve().parents[2]
    / "fixtures"
    / "qre-samples"
    / "S01_campus_cafeteria_experience.docx"
)


def _build(path: Path, headings: list[tuple[str, str]], preamble: str | None = None):
    """Write a .docx with the given (style, text) blocks, one body para each."""
    d = docx.Document()
    if preamble:
        d.add_paragraph(preamble)
    for style, text in headings:
        p = d.add_paragraph(text)
        p.style = d.styles[style]
        d.add_paragraph(f"body content under {text}")
    d.save(str(path))


def _sections_of(tmp: str, headings, preamble=None, **kwargs):
    path = Path(tmp) / "s.docx"
    _build(path, headings, preamble)
    return detect_sections(read_docx(path), **kwargs)


def test_known_headings_classified_as_extracted():
    with tempfile.TemporaryDirectory() as tmp:
        sd = _sections_of(
            tmp,
            [("Heading 1", "Study specification"), ("Heading 1", "Questionnaire")],
        )
        labels = [(s.label, s.label_provenance) for s in sd.sections]
        assert ("study_specification", LABEL_EXTRACTED) in labels
        assert ("questionnaire", LABEL_EXTRACTED) in labels


def test_heading_aliases_and_punctuation_normalized():
    """'Routing & Termination:' must resolve to the same label as the canonical name."""
    with tempfile.TemporaryDirectory() as tmp:
        sd = _sections_of(tmp, [("Heading 2", "Routing & Termination:")])
        section = next(s for s in sd.sections if s.heading_text.startswith("Routing"))
        assert section.label == "routing_and_termination"
        assert section.label_provenance == LABEL_EXTRACTED
        assert section.requires_review is False


def test_unknown_heading_is_flagged_not_guessed():
    with tempfile.TemporaryDirectory() as tmp:
        sd = _sections_of(tmp, [("Heading 1", "Weighting And Analysis Plan")])
        section = next(s for s in sd.sections if s.heading_text.startswith("Weighting"))
        assert section.label is None
        assert section.label_provenance == LABEL_UNKNOWN
        assert section.requires_review is True
        # Content is preserved despite being unclassified (CLAUDE.md §16)
        assert len(section.blocks) == 1
        assert any(r.reason == "heading_not_classified" for r in sd.review_queue)
        # ...and it is excluded from the flat dict view, which has no key for it
        assert sd.unclassified_sections == [section]


def test_bold_body_text_is_not_treated_as_a_heading():
    """The study-spec lines in the real QRE are bold body text, not headings."""
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "bold.docx"
        d = docx.Document()
        h = d.add_paragraph("Study specification")
        h.style = d.styles["Heading 1"]
        for line in ("Business objective: x", "Mode: web", "Estimated length: 4 min"):
            p = d.add_paragraph()
            p.add_run(line).bold = True
        d.save(str(path))

        sd = detect_sections(read_docx(path))
        spec = next(s for s in sd.sections if s.label == "study_specification")
        # One section holding three bold lines — not three sections
        assert len(spec.blocks) == 3
        assert all(b.is_bold for b in spec.blocks)


def test_preamble_captured_before_first_heading():
    with tempfile.TemporaryDirectory() as tmp:
        sd = _sections_of(
            tmp, [("Heading 1", "Questionnaire")], preamble="QRE title page line"
        )
        first = sd.sections[0]
        assert first.label == PREAMBLE_LABEL
        assert first.label_provenance == LABEL_DERIVED
        assert first.heading_text == ""
        assert "title page" in first.blocks[0].text


def test_document_with_no_headings_is_flagged():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "flat.docx"
        d = docx.Document()
        d.add_paragraph("Just body text, no styled headings at all.")
        d.save(str(path))

        sd = detect_sections(read_docx(path))
        assert any(r.reason == "no_headings_detected" for r in sd.review_queue)
        # Content still reachable rather than dropped
        assert sd.sections[0].label == PREAMBLE_LABEL
        assert len(sd.sections[0].blocks) == 1


def test_injected_classifier_used_only_for_unknown_headings():
    calls: list[str] = []

    def classifier(heading_text, allowed):
        calls.append(heading_text)
        return "quota_controls"

    with tempfile.TemporaryDirectory() as tmp:
        sd = _sections_of(
            tmp,
            [("Heading 1", "Questionnaire"), ("Heading 1", "Sample Balancing Rules")],
            classifier=classifier,
        )
        # Known heading resolved deterministically — classifier never consulted
        assert calls == ["Sample Balancing Rules"]
        inferred = next(s for s in sd.sections if s.heading_text.startswith("Sample"))
        assert inferred.label == "quota_controls"
        assert inferred.label_provenance == LABEL_INFERRED


def test_classifier_cannot_invent_labels_outside_vocabulary():
    """A classifier returning an unknown label must not widen the contract (§17)."""
    with tempfile.TemporaryDirectory() as tmp:
        sd = _sections_of(
            tmp,
            [("Heading 1", "Totally Novel Section")],
            classifier=lambda text, allowed: "some_made_up_label",
        )
        section = sd.sections[-1]
        assert section.label is None
        assert section.label_provenance == LABEL_UNKNOWN


def test_custom_vocabulary_overrides_default():
    with tempfile.TemporaryDirectory() as tmp:
        sd = _sections_of(
            tmp,
            [("Heading 1", "Fieldwork Notes")],
            section_types={"fieldwork": ("fieldwork notes",)},
        )
        assert sd.sections[-1].label == "fieldwork"
        assert sd.known_labels == ["fieldwork"]


def test_no_block_is_lost_or_duplicated():
    """Every input block lands in exactly one section, or is a heading."""
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "s.docx"
        _build(
            path,
            [("Heading 1", "Questionnaire"), ("Heading 1", "Unknown Thing")],
            preamble="preamble line",
        )
        doc = read_docx(path)
        sd = detect_sections(doc)

        in_sections = [b.order_index for s in sd.sections for b in s.blocks]
        headings = [
            b.order_index
            for b in doc.blocks
            if b.kind == "paragraph" and b.style_name.lower().startswith("heading")
        ]
        assert len(in_sections) == len(set(in_sections)), "a block was duplicated"
        assert sorted(in_sections + headings) == [b.order_index for b in doc.blocks]


def test_real_fixture_if_present():
    """Sanity check against the actual sample QRE, when available locally."""
    if not FIXTURE.exists():
        return
    sd = detect_sections(read_docx(FIXTURE))
    by_label = sd.by_label()

    for expected in (
        "study_specification",
        "questionnaire",
        "routing_and_termination",
        "programming_and_qa_requirements",
        "acceptance_test_scenarios",
        "completion_messages",
    ):
        assert expected in by_label, f"missing section: {expected}"

    # Questionnaire and routing sections must carry their tables through as grids
    assert any(b.kind == "table" for b in by_label["questionnaire"])
    assert any(b.kind == "table" for b in by_label["routing_and_termination"])

    # Nothing in this document should be unclassified
    assert sd.unclassified_sections == []

    # S01 has no quota section; that must be reported, not invented
    assert "quota_controls" not in by_label
    assert any(r.reason == "known_section_not_found" for r in sd.review_queue)


if __name__ == "__main__":
    for name, fn in sorted(list(globals().items())):
        if name.startswith("test_") and callable(fn):
            fn()
    print(f"OK: all section_detector checks passed ({len(DEFAULT_SECTION_TYPES)} known types)")
