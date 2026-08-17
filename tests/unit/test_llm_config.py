"""Checks for config loading and the Groq classification prompt.

No network calls: the client is replaced with a stub throughout. Tests that a
real key is never required, never logged, and that every failure path degrades to
None so the caller flags rather than guesses.

Run directly: python3 tests/unit/test_llm_config.py
Run via pytest: python3 -m pytest tests/unit/test_llm_config.py -v
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "src"))

from common.config import GroqSettings, Settings, _parse_env_file, get_settings
from common.llm.groq_client import GroqClient, LLMCallError, LLMUnavailableError, build_client
from common.prompts.qre_interpretation import classify_section_heading

ALLOWED = ["questionnaire", "quota_controls", "routing_and_termination"]


class StubClient:
    """Stands in for GroqClient. Records prompts, returns a scripted payload."""

    def __init__(self, payload=None, error: Exception | None = None):
        self._payload = payload if payload is not None else {"label": None}
        self._error = error
        self.calls: list[tuple[str, str]] = []
        self.model = "stub-model"

    def complete_json(self, system_prompt, user_prompt, *, temperature=0.0):
        self.calls.append((system_prompt, user_prompt))
        if self._error is not None:
            raise self._error
        return self._payload


# ---------------------------------------------------------------------------
# .env parsing
# ---------------------------------------------------------------------------


def test_env_file_parsing_handles_comments_blanks_and_quotes():
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / ".env"
        path.write_text(
            "# a comment\n"
            "\n"
            "GROQ_API_KEY=abc123\n"
            'GROQ_MODEL="quoted-model"\n'
            "GROQ_TIMEOUT_SECONDS = 45 \n"
            "MALFORMED_LINE_NO_EQUALS\n"
        )
        values = _parse_env_file(path)
        assert values["GROQ_API_KEY"] == "abc123"
        assert values["GROQ_MODEL"] == "quoted-model"
        assert values["GROQ_TIMEOUT_SECONDS"] == "45"
        assert "MALFORMED_LINE_NO_EQUALS" not in values


def test_missing_env_file_is_not_an_error():
    assert _parse_env_file(Path("/nonexistent/.env")) == {}


def test_real_environment_overrides_env_file(monkeypatch=None):
    """An exported variable must win over the .env entry."""
    import os

    original = os.environ.get("GROQ_MODEL")
    try:
        os.environ["GROQ_MODEL"] = "from-environment"
        get_settings.cache_clear()
        assert get_settings().groq.model == "from-environment"
    finally:
        if original is None:
            os.environ.pop("GROQ_MODEL", None)
        else:
            os.environ["GROQ_MODEL"] = original
        get_settings.cache_clear()


# ---------------------------------------------------------------------------
# Settings behaviour
# ---------------------------------------------------------------------------


def _settings(provider: str, key: str) -> Settings:
    return Settings(
        llm_provider=provider,
        groq=GroqSettings(api_key=key, model="m", timeout_seconds=5, max_retries=0),
    )


def test_llm_disabled_without_key():
    assert _settings("groq", "").llm_enabled is False


def test_llm_disabled_when_provider_none_even_with_key():
    assert _settings("none", "real-key").llm_enabled is False


def test_llm_enabled_only_with_provider_and_key():
    assert _settings("groq", "real-key").llm_enabled is True


def test_api_key_never_appears_in_repr():
    """A key must not leak into logs, tracebacks or debugger output (§51, §55)."""
    secret = "gsk_supersecretvalue_do_not_leak"
    settings = GroqSettings(
        api_key=secret, model="m", timeout_seconds=5, max_retries=0
    )
    assert secret not in repr(settings)
    assert secret not in str(settings)
    assert "<set>" in repr(settings)
    # And not via the enclosing Settings object either
    assert secret not in repr(_settings("groq", secret))


# ---------------------------------------------------------------------------
# Client construction
# ---------------------------------------------------------------------------


def test_client_refuses_to_build_without_key():
    try:
        GroqClient(GroqSettings(api_key="", model="m", timeout_seconds=5, max_retries=0))
        raise AssertionError("expected LLMUnavailableError")
    except LLMUnavailableError as exc:
        assert "GROQ_API_KEY" in str(exc)


def test_build_client_returns_none_instead_of_raising():
    """Callers that must degrade gracefully get None, not an exception."""
    settings = GroqSettings(api_key="", model="m", timeout_seconds=5, max_retries=0)
    assert build_client(settings) is None


# ---------------------------------------------------------------------------
# Step 2 classification prompt
# ---------------------------------------------------------------------------


def test_classifier_returns_label_from_payload():
    stub = StubClient({"label": "quota_controls"})
    assert classify_section_heading("Sample Balancing", ALLOWED, client=stub) == "quota_controls"
    assert len(stub.calls) == 1


def test_classifier_sends_only_heading_and_labels():
    """No document body may reach the provider (confidentiality)."""
    stub = StubClient({"label": "questionnaire"})
    classify_section_heading("Question List", ALLOWED, client=stub)
    _system, user = stub.calls[0]
    assert "Question List" in user
    for label in ALLOWED:
        assert label in user
    # The prompt is built from the heading and labels alone; nothing else is added.
    assert len(user) < 400


def test_classifier_returns_none_on_null_label():
    """An honest 'no match' from the model must stay a non-answer."""
    assert classify_section_heading("Weird Heading", ALLOWED, client=StubClient({"label": None})) is None


def test_classifier_returns_none_on_missing_or_blank_label():
    assert classify_section_heading("H", ALLOWED, client=StubClient({})) is None
    assert classify_section_heading("H", ALLOWED, client=StubClient({"label": "   "})) is None
    assert classify_section_heading("H", ALLOWED, client=StubClient({"label": 42})) is None


def test_classifier_returns_none_on_call_failure():
    """A provider outage must not stop the run (§30, §31)."""
    import logging

    stub = StubClient(error=LLMCallError("upstream 503"))
    # The warning it logs is expected here; silence it so it is not mistaken for
    # a real failure in test output.
    prompt_logger = logging.getLogger("common.prompts.qre_interpretation")
    prompt_logger.disabled = True
    try:
        assert classify_section_heading("Anything", ALLOWED, client=stub) is None
    finally:
        prompt_logger.disabled = False


def test_classifier_skips_call_for_blank_heading():
    stub = StubClient({"label": "questionnaire"})
    assert classify_section_heading("   ", ALLOWED, client=stub) is None
    assert stub.calls == [], "no request should be made for a blank heading"


def test_classifier_output_is_still_vocabulary_checked_by_step2():
    """The prompt may return anything; Step 2 is the gate that rejects it (§17)."""
    import tempfile as _tf

    import docx

    from agents.qre_interpretation.extraction.section_detector import detect_sections
    from agents.qre_interpretation.extraction.sectioned_document import LABEL_UNKNOWN
    from agents.qre_interpretation.ingestion.docx_reader import read_docx

    stub = StubClient({"label": "a_label_that_does_not_exist"})

    with _tf.TemporaryDirectory() as tmp:
        path = Path(tmp) / "s.docx"
        d = docx.Document()
        h = d.add_paragraph("Totally Novel Heading")
        h.style = d.styles["Heading 1"]
        d.add_paragraph("body")
        d.save(str(path))

        sectioned = detect_sections(
            read_docx(path),
            classifier=lambda text, allowed: classify_section_heading(
                text, allowed, client=stub
            ),
        )
        section = sectioned.sections[-1]
        assert section.label is None
        assert section.label_provenance == LABEL_UNKNOWN


if __name__ == "__main__":
    for _name, _fn in sorted(list(globals().items())):
        if _name.startswith("test_") and callable(_fn):
            _fn()
    print("OK: all llm/config checks passed")
